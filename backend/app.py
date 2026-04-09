# Main app runner for pitcrew
# Run with: uvicorn backend.app:app --reload
# FastAPI: REST endpoints, serves static/ as the frontend

import base64
import io
import json
import os
import time
import asyncio
from contextlib import asynccontextmanager
from pathlib import Path
from typing import Optional

from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from PIL import Image
from pydantic import BaseModel

from dotenv import load_dotenv
from google import genai
from google.genai import types

# Resolve .env relative to this file so it works regardless of cwd
load_dotenv(os.path.join(os.path.dirname(__file__), '..', '.env'))

from backend.db import (
    init_db,
    get_cars, get_car, create_car, update_car, delete_car, update_car_photo,
    get_journals, get_journal, create_journal, update_journal, delete_journal,
    get_views, get_view, create_view, delete_view,
    get_photopins, get_photopin, create_photopin, update_photopin, delete_photopin,
    get_parts, get_part, create_part, update_part, delete_part,
    get_manuals, get_manual, create_manual, delete_manual,
)

# from backend.prompt ()
# from backend.voice ()

# Directories for static files and uploads
STATIC_DIR = os.path.join(os.path.dirname(__file__), '..', 'static')
UPLOADS_DIR = os.path.join(STATIC_DIR, 'uploads')
MANUALS_DIR = os.path.join(UPLOADS_DIR, 'manuals')
VIEWS_DIR = os.path.join(UPLOADS_DIR, 'views')
os.makedirs(UPLOADS_DIR, exist_ok=True)
os.makedirs(MANUALS_DIR, exist_ok=True)
os.makedirs(VIEWS_DIR, exist_ok=True)

gemini = genai.Client(api_key=os.getenv('GENAI_API_KEY'))


# ── Document helpers ─────────────────────────────────────────────────────────

def _extract_sidecar(filepath: str) -> None:
    """Extract text from a PDF or DOCX and save as a .txt sidecar file. Raises on failure."""
    ext = Path(filepath).suffix.lower()
    sidecar = filepath + '.txt'
    if ext == '.pdf':
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ''
            if text.strip():
                pages.append(f'[Page {i + 1}]\n{text}')
        with open(sidecar, 'w', encoding='utf-8') as f:
            f.write('\n\n'.join(pages))
    elif ext in ('.docx', '.doc'):
        import docx as docx_lib
        doc = docx_lib.Document(filepath)
        text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
        with open(sidecar, 'w', encoding='utf-8') as f:
            f.write(text)


def _relevant_chunks(text: str, question: str, chunk_size: int = 8000, top_n: int = 6) -> list[str]:
    """Return the top_n chunks of text most relevant to the question by keyword overlap."""
    chunks = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]
    if not chunks:
        return []
    stop = {
        'the', 'a', 'an', 'is', 'are', 'was', 'were', 'be', 'been', 'have', 'has',
        'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may',
        'might', 'can', 'for', 'of', 'to', 'in', 'on', 'at', 'by', 'from',
        'with', 'and', 'or', 'but', 'if', 'what', 'how', 'where', 'when',
        'which', 'who', 'that', 'this', 'i', 'my', 'me',
    }
    keywords = {w for w in question.lower().split() if w not in stop and len(w) > 2}
    scored = sorted(
        ((sum(1 for kw in keywords if kw in chunk.lower()), chunk) for chunk in chunks),
        key=lambda x: x[0], reverse=True,
    )
    result = [c for score, c in scored[:top_n] if score > 0]
    return result or [chunks[0]]


# Lifespan hooks for startup/shutdown tasks like DB initialization, cleanup, etc.
@asynccontextmanager
async def lifespan(app: FastAPI):
    await init_db()
    yield

app = FastAPI(lifespan=lifespan)

# ── Static frontend ───────────────────────────────────────────────────────────

app.mount('/static', StaticFiles(directory=STATIC_DIR), name='static')

@app.get('/')
async def index():
    return FileResponse(os.path.join(STATIC_DIR, 'index.html'))

# ── Cars ──────────────────────────────────────────────────────────────────────

class CarBody(BaseModel):
    year: Optional[int] = None
    make: Optional[str] = None
    model: Optional[str] = None
    trim: Optional[str] = None
    options: Optional[str] = None
    engine: Optional[str] = None
    color: Optional[str] = None
    vin: Optional[str] = None
    notes: Optional[str] = None


class CarPatch(BaseModel):
    options: Optional[str] = None
    year: Optional[int] = None
    make: Optional[str] = None
    model: Optional[str] = None
    trim: Optional[str] = None
    engine: Optional[str] = None
    color: Optional[str] = None
    vin: Optional[str] = None
    notes: Optional[str] = None


@app.get('/api/cars')
async def list_cars():
    return await get_cars()


@app.post('/api/cars', status_code=201)
async def add_car(body: CarBody):
    car_id = await create_car(body.model_dump())
    return await get_car(car_id)


@app.get('/api/cars/{car_id}')
async def fetch_car(car_id: int):
    car = await get_car(car_id)
    if not car:
        raise HTTPException(404, 'Car not found')
    return car


@app.patch('/api/cars/{car_id}')
async def edit_car(car_id: int, body: CarPatch):
    car = await get_car(car_id)
    if not car:
        raise HTTPException(404, 'Car not found')
    await update_car(car_id, body.model_dump(exclude_none=True))
    return await get_car(car_id)


@app.delete('/api/cars/{car_id}', status_code=204)
async def remove_car(car_id: int):
    car = await get_car(car_id)
    if not car:
        raise HTTPException(404, 'Car not found')
    await delete_car(car_id)


@app.post('/api/cars/{car_id}/photo')
async def upload_car_photo(car_id: int, file: UploadFile = File(...)):
    car = await get_car(car_id)
    if not car:
        raise HTTPException(404, 'Car not found')
    ext = Path(file.filename).suffix.lower() if file.filename else '.jpg'
    if ext not in ('.jpg', '.jpeg', '.png', '.webp', '.gif'):
        raise HTTPException(400, 'Unsupported file type')
    dest = os.path.join(UPLOADS_DIR, f'car_{car_id}{ext}')
    contents = await file.read()
    with open(dest, 'wb') as f:
        f.write(contents)
    photo_url = f'/static/uploads/car_{car_id}{ext}'
    await update_car_photo(car_id, photo_url)
    return {'photo_url': photo_url}


@app.delete('/api/cars/{car_id}/photo', status_code=204)
async def delete_car_photo(car_id: int):
    car = await get_car(car_id)
    if not car:
        raise HTTPException(404, 'Car not found')
    # Delete file if it exists
    if car.get('photo_url'):
        filename = car['photo_url'].split('/')[-1]
        filepath = os.path.join(UPLOADS_DIR, filename)
        if os.path.exists(filepath):
            os.remove(filepath)
    # Clear photo_url in DB
    await update_car_photo(car_id, None)


# ── Journal ───────────────────────────────────────────────────────────────────
class JournalEntryBody(BaseModel):
    type: str 
    title: str
    body: Optional[str] = None


class JournalEntryPatch(BaseModel):
    title: Optional[str] = None
    body: Optional[str] = None

@app.get('/api/cars/{car_id}/journal')
async def list_journals(car_id: int):
    car = await get_car(car_id)
    if not car:
        raise HTTPException(404, 'Car not found')
    return await get_journals(car_id)

@app.post('/api/cars/{car_id}/journal', status_code=201)
async def add_journal(car_id: int, body: JournalEntryBody):
    car = await get_car(car_id)
    if not car:
        raise HTTPException(404, 'Car not found')
    entry_id = await create_journal(car_id, body.model_dump())
    return await get_journal(entry_id)


@app.patch('/api/journal/{entry_id}')
async def edit_journal(entry_id: int, body: JournalEntryPatch):
    entry = await get_journal(entry_id)
    if not entry:
        raise HTTPException(404, 'Journal entry not found')
    await update_journal(entry_id, body.model_dump(exclude_none=True))
    return await get_journal(entry_id)


@app.delete('/api/journal/{entry_id}', status_code=204)
async def remove_journal(entry_id: int):
    entry = await get_journal(entry_id)
    if not entry:
        raise HTTPException(404, 'Journal entry not found')
    await delete_journal(entry_id)


class ResearchRequest(BaseModel):
    query: str


@app.post('/api/cars/{car_id}/research', status_code=201)
async def research(car_id: int, req: ResearchRequest):
    car = await get_car(car_id)
    if not car:
        raise HTTPException(404, 'Car not found')

    car_desc = ' '.join(filter(None, [
        str(car.get('year') or ''),
        car.get('make'),
        car.get('model'),
        car.get('trim'),
        f"({car['engine']})" if car.get('engine') else None,
    ]))
    options_line = f"\nOptions/package: {car['options']}" if car.get('options') else ''

    prompt = (
        f"You are an expert automotive mechanic and technical researcher.\n"
        f"The car in question is: {car_desc}.{options_line}\n\n"
        f"Answer the following question concisely and practically. "
        f"Focus on specs, procedures, torque values, part numbers, and facts "
        f"specific to this car. Keep your answer under 400 words.\n\n"
        f"Question: {req.query}"
    )

    try:
        response = await gemini.aio.models.generate_content(
            model='gemini-flash-latest',
            contents=prompt,
            config=types.GenerateContentConfig(temperature=0.3),
        )
        answer = response.text
    except Exception as e:
        raise HTTPException(502, f'AI request failed: {e}')

    entry_id = await create_journal(car_id, {
        'type': 'research',
        'title': req.query,
        'body': answer,
    })
    return await get_journal(entry_id)


# ── Photo Pins ─────────────────────────────────────────────────────────────────

class PhotoPinBody(BaseModel):
    label: str
    notes: Optional[str] = None
    x_pct: Optional[float] = None
    y_pct: Optional[float] = None


class PhotoPinPatch(BaseModel):
    label: Optional[str] = None
    notes: Optional[str] = None
    x_pct: Optional[float] = None
    y_pct: Optional[float] = None


@app.get('/api/cars/{car_id}/views')
async def list_views(car_id: int):
    car = await get_car(car_id)
    if not car:
        raise HTTPException(404, 'Car not found')
    return await get_views(car_id)


@app.post('/api/cars/{car_id}/views/{angle}')
async def upload_view(car_id: int, angle: str, file: UploadFile = File(...)):
    valid_angles = ('front', 'sideD', 'sideP', 'rear', 'engine', 'underside', 'interior')
    if angle not in valid_angles:
        raise HTTPException(400, f'Angle must be one of {valid_angles}')
    car = await get_car(car_id)
    if not car:
        raise HTTPException(404, 'Car not found')
    ext = Path(file.filename).suffix.lower() if file.filename else '.jpg'
    if ext not in ('.jpg', '.jpeg', '.png', '.webp'):
        raise HTTPException(400, 'Unsupported file type')
    # Millisecond timestamp ensures multiple uploads per angle never collide
    filename = f'view_{car_id}_{angle}_{int(time.time() * 1000)}{ext}'
    dest = os.path.join(VIEWS_DIR, filename)
    contents = await file.read()
    with open(dest, 'wb') as f:
        f.write(contents)
    file_path = f'/static/uploads/views/{filename}'
    name = Path(file.filename).stem if file.filename else angle
    view_id = await create_view(car_id, name, angle, file_path)
    return await get_view(view_id)


@app.delete('/api/views/{view_id}', status_code=204)
async def remove_view(view_id: int):
    view = await get_view(view_id)
    if not view:
        raise HTTPException(404, 'View not found')
    if view.get('file_path'):
        filename = view['file_path'].split('/')[-1]
        filepath = os.path.join(VIEWS_DIR, filename)
        if os.path.exists(filepath):
            os.remove(filepath)
    await delete_view(view_id)


@app.get('/api/views/{view_id}/pins')
async def list_pins(view_id: int):
    view = await get_view(view_id)
    if not view:
        raise HTTPException(404, 'View not found')
    return await get_photopins(view_id)


@app.post('/api/views/{view_id}/pins', status_code=201)
async def add_pin(view_id: int, body: PhotoPinBody):
    view = await get_view(view_id)
    if not view:
        raise HTTPException(404, 'View not found')
    pin_id = await create_photopin(view_id, body.model_dump())
    return await get_photopin(pin_id)


@app.post('/api/pins/{pin_id}/research')
async def research_pin(pin_id: int):
    """Send photo + pin context to Gemini and return research about that specific component."""
    pin = await get_photopin(pin_id)
    if not pin:
        raise HTTPException(404, 'Pin not found')

    view = await get_view(pin['view_id'])
    if not view:
        raise HTTPException(404, 'View not found')

    car = await get_car(view['car_id'])
    car_desc = ' '.join(filter(None, [
        str(car.get('year') or ''),
        car.get('make'),
        car.get('model'),
        car.get('trim'),
        f"({car['engine']})" if car.get('engine') else None,
    ])) if car else 'Unknown car'

    angle_labels = {
        'front':     'front view',
        'sideD':     "driver's side view",
        'sideP':     "passenger's side view",
        'rear':      'rear view',
        'engine':    'engine bay view',
        'underside': 'underside / chassis view',
        'interior':  'interior / cabin view',
    }
    angle_label = angle_labels.get(view.get('angle', ''), view.get('angle', 'view'))

    relative = view['file_path'].removeprefix('/static/')
    filepath = os.path.join(STATIC_DIR, relative)
    if not os.path.exists(filepath):
        raise HTTPException(404, 'Image file not found on disk')

    with open(filepath, 'rb') as f:
        raw_bytes = f.read()

    # Strip EXIF before sending to Gemini
    img = Image.open(io.BytesIO(raw_bytes))
    fmt = (img.format or 'JPEG').upper()
    buf = io.BytesIO()
    save_kwargs: dict = {'format': fmt}
    if fmt in ('JPEG', 'WEBP'):
        save_kwargs['exif'] = b''
    img.save(buf, **save_kwargs)
    clean_bytes = buf.getvalue()

    mime_map = {
        'JPEG': 'image/jpeg', 'PNG': 'image/png',
        'WEBP': 'image/webp', 'GIF': 'image/gif',
    }
    mime_type = mime_map.get(fmt, 'image/jpeg')

    label = pin.get('label', '')
    x_pct = pin.get('x_pct')
    y_pct = pin.get('y_pct')
    pin_notes = pin.get('notes', '')

    position_hint = ''
    if x_pct is not None and y_pct is not None:
        position_hint = (
            f"The pin is placed at approximately {x_pct:.0f}% from the left "
            f"and {y_pct:.0f}% from the top of the image."
        )

    prompt_text = (
        f"You are an expert automotive mechanic and technical researcher.\n"
        f"The vehicle is: {car_desc}. This photo is the {angle_label}.\n\n"
        f"A pin has been placed on a specific component labelled: \"{label}\".\n"
        f"{position_hint}\n"
        f"{f'Additional context from the user: {pin_notes}' if pin_notes else ''}\n\n"
        f"Please research this specific component and provide:\n"
        f"1. Exact component name and function\n"
        f"2. OEM part number(s) for this vehicle\n"
        f"3. Common aftermarket / compatible alternatives\n"
        f"4. Service interval or maintenance notes\n"
        f"5. Known failure modes or things to watch for\n\n"
        f"Be specific to this vehicle. Keep the answer under 500 words. "
        f"Use the photo as visual context to confirm the component if visible."
    )

    try:
        response = await gemini.aio.models.generate_content(
            model='gemini-flash-latest',
            contents=types.Content(
                role='user',
                parts=[
                    types.Part.from_bytes(data=clean_bytes, mime_type=mime_type),
                    types.Part.from_text(text=prompt_text),
                ]
            ),
            config=types.GenerateContentConfig(temperature=0.2),
        )
        summary = response.text.strip()
        await update_photopin(pin_id, {'ai_summary': summary})
        return {'summary': summary}
    except Exception as e:
        raise HTTPException(502, f'AI research failed: {e}')


@app.delete('/api/pins/{pin_id}', status_code=204)
async def remove_pin(pin_id: int):
    pin = await get_photopin(pin_id)
    if not pin:
        raise HTTPException(404, 'Pin not found')
    await delete_photopin(pin_id)


# ── Parts ─────────────────────────────────────────────────────────────────────

class PartBody(BaseModel):
    name: str
    part_number: Optional[str] = None
    supplier: Optional[str] = None
    url: Optional[str] = None
    price: Optional[float] = None
    quantity: Optional[int] = 1
    category: Optional[str] = 'Mechanical'
    notes: Optional[str] = None


class PartPatch(BaseModel):
    name: Optional[str] = None
    part_number: Optional[str] = None
    supplier: Optional[str] = None
    url: Optional[str] = None
    price: Optional[float] = None
    quantity: Optional[int] = None
    category: Optional[str] = None
    status: Optional[str] = None
    notes: Optional[str] = None


@app.get('/api/cars/{car_id}/parts')
async def list_parts(car_id: int):
    car = await get_car(car_id)
    if not car:
        raise HTTPException(404, 'Car not found')
    return await get_parts(car_id)


@app.post('/api/cars/{car_id}/parts', status_code=201)
async def add_part(car_id: int, body: PartBody):
    car = await get_car(car_id)
    if not car:
        raise HTTPException(404, 'Car not found')
    part_id = await create_part(car_id, body.model_dump())
    return await get_part(part_id)


@app.patch('/api/parts/{part_id}')
async def edit_part(part_id: int, body: PartPatch):
    part = await get_part(part_id)
    if not part:
        raise HTTPException(404, 'Part not found')
    await update_part(part_id, body.model_dump(exclude_none=True))
    return await get_part(part_id)


@app.delete('/api/parts/{part_id}', status_code=204)
async def remove_part(part_id: int):
    part = await get_part(part_id)
    if not part:
        raise HTTPException(404, 'Part not found')
    await delete_part(part_id)


# ── Manuals ───────────────────────────────────────────────────────────────────

@app.get('/api/cars/{car_id}/manuals')
async def list_manuals(car_id: int):
    car = await get_car(car_id)
    if not car:
        raise HTTPException(404, 'Car not found')
    return await get_manuals(car_id)


@app.post('/api/cars/{car_id}/manuals', status_code=201)
async def upload_manual(car_id: int, file: UploadFile = File(...), category: str = 'manual'):
    car = await get_car(car_id)
    if not car:
        raise HTTPException(404, 'Car not found')
    if category not in ('manual', 'reference', 'other'):
        raise HTTPException(400, "category must be 'manual', 'reference', or 'other'")
    ALLOWED: dict = {
        'manual':    {'.pdf', '.docx', '.doc'},
        'reference': {'.jpg', '.jpeg', '.png', '.webp', '.gif', '.svg', '.bmp'},
        'other':     None,  # any file type accepted
    }
    ext = Path(file.filename).suffix.lower() if file.filename else ''
    allowed = ALLOWED[category]
    if allowed is not None and ext not in allowed:
        raise HTTPException(400, f'Unsupported file type for {category}. Allowed: {", ".join(sorted(allowed))}')
    safe_name = Path(file.filename).stem[:60] if file.filename else 'file'
    filename = f"car_{car_id}_{int(time.time())}_{safe_name}{ext}"
    dest = os.path.join(MANUALS_DIR, filename)
    contents = await file.read()
    with open(dest, 'wb') as f:
        f.write(contents)
    if ext in ('.pdf', '.docx', '.doc'):
        await asyncio.to_thread(_extract_sidecar, dest)
    file_path = f'/static/uploads/manuals/{filename}'
    title = Path(file.filename).stem if file.filename else filename
    manual_id = await create_manual(car_id, title, file_path, category)
    return await get_manual(manual_id)


@app.delete('/api/manuals/{manual_id}', status_code=204)
async def remove_manual(manual_id: int):
    manual = await get_manual(manual_id)
    if not manual:
        raise HTTPException(404, 'Manual not found')
    # Delete file from disk
    if manual.get('file_path'):
        filename = manual['file_path'].split('/')[-1]
        filepath = os.path.join(MANUALS_DIR, filename)
        if os.path.exists(filepath):
            os.remove(filepath)
        sidecar = filepath + '.txt'
        if os.path.exists(sidecar):
            os.remove(sidecar)
    await delete_manual(manual_id)


class ManualAskBody(BaseModel):
    question: str
    manual_ids: list[int]


@app.post('/api/cars/{car_id}/manuals/ask')
async def ask_manuals(car_id: int, body: ManualAskBody):
    """Answer a question using only the selected uploaded documents."""
    car = await get_car(car_id)
    if not car:
        raise HTTPException(404, 'Car not found')
    if not body.manual_ids:
        raise HTTPException(400, 'Select at least one document')

    car_desc = ' '.join(filter(None, [
        str(car.get('year') or ''), car.get('make'), car.get('model'), car.get('trim'),
    ])) or 'Unknown car'

    mime_map = {
        '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg', '.png': 'image/png',
        '.webp': 'image/webp', '.gif': 'image/gif', '.bmp': 'image/bmp',
        '.svg': 'image/svg+xml',
    }

    content_parts: list = []
    loaded: list[str] = []

    for mid in body.manual_ids:
        doc = await get_manual(mid)
        if not doc:
            continue
        fp = doc.get('file_path', '')
        relative = fp.removeprefix('/static/')
        filepath = os.path.join(STATIC_DIR, relative)
        if not os.path.exists(filepath):
            continue
        ext = Path(filepath).suffix.lower()
        try:
            if ext == '.pdf':
                sidecar = filepath + '.txt'
                if not os.path.exists(sidecar):
                    try:
                        await asyncio.to_thread(_extract_sidecar, filepath)
                    except Exception as extract_err:
                        raise HTTPException(502, f'Could not extract text from "{doc["title"]}": {extract_err}. '
                                                 f'It may be a scanned/image-only PDF.')
                with open(sidecar, 'r', encoding='utf-8') as f:
                    full_text = f.read().strip()
                if not full_text:
                    raise HTTPException(400, f'"{doc["title"]}" appears to be a scanned PDF with no '
                                            f'extractable text. AI search requires a text-based PDF.')
                chunks = _relevant_chunks(full_text, body.question)
                excerpt = '\n\n---\n\n'.join(chunks)
                content_parts.append(types.Part.from_text(
                    text=f'[Relevant excerpts from: {doc["title"]}]\n\n{excerpt}'
                ))
            elif ext in mime_map:
                with open(filepath, 'rb') as f:
                    content_parts.append(types.Part.from_bytes(data=f.read(), mime_type=mime_map[ext]))
                content_parts.append(types.Part.from_text(text=f'[Above image: {doc["title"]}]'))
            elif ext in ('.docx', '.doc'):
                sidecar = filepath + '.txt'
                if os.path.exists(sidecar):
                    with open(sidecar, 'r', encoding='utf-8') as f:
                        full_text = f.read()
                else:
                    import docx as docx_lib
                    document = docx_lib.Document(filepath)
                    full_text = '\n'.join(p.text for p in document.paragraphs if p.text.strip())
                chunks = _relevant_chunks(full_text, body.question)
                content_parts.append(types.Part.from_text(
                    text=f'[Relevant excerpts from: {doc["title"]}]\n\n' + '\n\n---\n\n'.join(chunks)
                ))
            else:
                with open(filepath, 'r', errors='replace') as f:
                    raw_text = f.read()
                chunks = _relevant_chunks(raw_text, body.question)
                content_parts.append(types.Part.from_text(
                    text=f'[File: {doc["title"]}]\n' + '\n---\n'.join(chunks)
                ))
            loaded.append(doc['title'])
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(502, f'Error reading "{doc.get("title", "document")}": {e}')

    if not content_parts:
        raise HTTPException(400, 'Could not read any of the selected documents')

    content_parts.append(types.Part.from_text(text=(
        f'You are an expert automotive technician assistant for a {car_desc}.\n'
        f'The content above consists of relevant excerpts from this car\'s uploaded documents.\n'
        f'Answer the following question as specifically as possible, citing page numbers or '
        f'section names from the excerpts where available.\n\n'
        f'Question: {body.question}'
    )))

    try:
        response = await gemini.aio.models.generate_content(
            model='gemini-flash-latest',
            contents=types.Content(role='user', parts=content_parts),
            config=types.GenerateContentConfig(temperature=0.1),
        )
        return {'answer': response.text.strip(), 'sources': loaded}
    except Exception as e:
        raise HTTPException(502, f'AI document search failed: {e}')



class ChatRequest(BaseModel):
    message: str
    car_id: Optional[int] = None


@app.post('/api/chat')
async def chat(req: ChatRequest):
    return {
        'response': f'(stub) you said: {req.message}',
        'citations': [],
    }
