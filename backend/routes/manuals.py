# Manual / Documents routes

import asyncio
import os
from pathlib import Path
import time

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from pydantic import BaseModel
from google.genai import types

from backend.auth import rate_limit_ai
from backend.db import (
    get_car, get_manual, get_manuals,
    create_manual, soft_delete_manual,
    VALID_MANUAL_CATEGORIES,
)
from backend.services import gemini

router = APIRouter(prefix='/api', tags=['manuals'])

STATIC_DIR = os.path.join(os.path.dirname(__file__), '..', '..', 'static')
MANUALS_DIR = os.path.join(STATIC_DIR, 'uploads', 'manuals')

MAX_MANUAL_UPLOAD = 100 * 1024 * 1024  # 100 MB

ALLOWED_EXTENSIONS: dict = {
    'manual':    {'.pdf', '.docx', '.doc'},
    'reference': {'.jpg', '.jpeg', '.png', '.webp', '.gif', '.svg', '.bmp'},
    'other':     None,  # any
}

MIME_MAP = {
    '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg', '.png': 'image/png',
    '.webp': 'image/webp', '.gif': 'image/gif', '.bmp': 'image/bmp',
    '.svg': 'image/svg+xml',
}


# ── Document text extraction ────────────────────────────────────────────────

def _extract_sidecar(filepath: str) -> None:
    """Extract text from a PDF or DOCX and save as a .txt sidecar file."""
    ext = Path(filepath).suffix.lower()
    sidecar = filepath + '.txt'
    if ext == '.pdf':
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ''
            if text.strip():
                pages.append(f'[Page {i + 1}]\n{text}')
        with open(sidecar, 'w', encoding='utf-8') as f:
            f.write('\n\n'.join(pages))
    elif ext in ('.docx', '.doc'):
        import docx as docx_lib
        doc = docx_lib.Document(filepath)
        text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
        with open(sidecar, 'w', encoding='utf-8') as f:
            f.write(text)


# ── Routes ───────────────────────────────────────────────────────────────────

@router.get('/cars/{car_id}/manuals')
async def list_manuals(car_id: int):
    car = await get_car(car_id)
    if not car:
        raise HTTPException(404, 'Car not found')
    return await get_manuals(car_id)


@router.post('/cars/{car_id}/manuals', status_code=201)
async def upload_manual(car_id: int, file: UploadFile = File(...), category: str = 'manual'):
    car = await get_car(car_id)
    if not car:
        raise HTTPException(404, 'Car not found')
    if category not in VALID_MANUAL_CATEGORIES:
        raise HTTPException(400, f"category must be one of {sorted(VALID_MANUAL_CATEGORIES)}")

    ext = Path(file.filename).suffix.lower() if file.filename else ''
    allowed = ALLOWED_EXTENSIONS[category]
    if allowed is not None and ext not in allowed:
        raise HTTPException(400, f'Unsupported file type for {category}. Allowed: {", ".join(sorted(allowed))}')

    contents = await file.read()
    if len(contents) > MAX_MANUAL_UPLOAD:
        raise HTTPException(400, f'File too large (max {MAX_MANUAL_UPLOAD // (1024*1024)}MB)')

    safe_name = Path(file.filename).stem[:60] if file.filename else 'file'
    filename = f"car_{car_id}_{int(time.time())}_{safe_name}{ext}"
    dest = os.path.join(MANUALS_DIR, filename)
    with open(dest, 'wb') as f:
        f.write(contents)

    # Extract text sidecar in background thread (non-blocking)
    if ext in ('.pdf', '.docx', '.doc'):
        await asyncio.to_thread(_extract_sidecar, dest)

    file_path = f'/static/uploads/manuals/{filename}'
    title = Path(file.filename).stem if file.filename else filename
    try:
        manual_id = await create_manual(car_id, title, file_path, category)
    except ValueError as e:
        os.remove(dest)
        raise HTTPException(400, str(e))
    return await get_manual(manual_id)


@router.delete('/manuals/{manual_id}', status_code=204)
async def remove_manual(manual_id: int):
    manual = await get_manual(manual_id)
    if not manual:
        raise HTTPException(404, 'Manual not found')
    # Soft delete — keep files on disk
    rows = await soft_delete_manual(manual_id)
    if rows == 0:
        raise HTTPException(404, 'Manual not found')


class ManualAskBody(BaseModel):
    question: str
    manual_ids: list[int]


@router.post('/cars/{car_id}/manuals/ask', dependencies=[Depends(rate_limit_ai)])
async def ask_manuals(car_id: int, body: ManualAskBody):
    """Answer a question using only the selected uploaded documents."""
    car = await get_car(car_id)
    if not car:
        raise HTTPException(404, 'Car not found')
    if not body.manual_ids:
        raise HTTPException(400, 'Select at least one document')

    content_parts: list = []
    loaded: list[str] = []

    for mid in body.manual_ids:
        doc = await get_manual(mid)
        if not doc:
            continue
        fp = doc.get('file_path', '')
        relative = fp.removeprefix('/static/')
        filepath = os.path.normpath(os.path.join(STATIC_DIR, relative))
        if not filepath.startswith(os.path.normpath(STATIC_DIR)):
            continue
        if not os.path.exists(filepath):
            continue
        ext = Path(filepath).suffix.lower()
        try:
            if ext == '.pdf':
                sidecar = filepath + '.txt'
                if not os.path.exists(sidecar):
                    try:
                        await asyncio.to_thread(_extract_sidecar, filepath)
                    except Exception as extract_err:
                        raise HTTPException(502, f'Could not extract text from "{doc["title"]}". '
                                                 f'It may be a scanned/image-only PDF.')
                with open(sidecar, 'r', encoding='utf-8') as f:
                    full_text = f.read().strip()
                if not full_text:
                    raise HTTPException(400, f'"{doc["title"]}" appears to be a scanned PDF with no '
                                            f'extractable text. AI search requires a text-based PDF.')
                chunks = gemini.relevant_chunks(full_text, body.question)
                excerpt = '\n\n---\n\n'.join(chunks)
                content_parts.append(types.Part.from_text(
                    text=f'[Relevant excerpts from: {doc["title"]}]\n\n{excerpt}'
                ))
            elif ext in MIME_MAP:
                with open(filepath, 'rb') as f:
                    content_parts.append(types.Part.from_bytes(data=f.read(), mime_type=MIME_MAP[ext]))
                content_parts.append(types.Part.from_text(text=f'[Above image: {doc["title"]}]'))
            elif ext in ('.docx', '.doc'):
                sidecar = filepath + '.txt'
                if os.path.exists(sidecar):
                    with open(sidecar, 'r', encoding='utf-8') as f:
                        full_text = f.read()
                else:
                    import docx as docx_lib
                    document = docx_lib.Document(filepath)
                    full_text = '\n'.join(p.text for p in document.paragraphs if p.text.strip())
                chunks = gemini.relevant_chunks(full_text, body.question)
                content_parts.append(types.Part.from_text(
                    text=f'[Relevant excerpts from: {doc["title"]}]\n\n' + '\n\n---\n\n'.join(chunks)
                ))
            else:
                with open(filepath, 'r', errors='replace') as f:
                    raw_text = f.read()
                chunks = gemini.relevant_chunks(raw_text, body.question)
                content_parts.append(types.Part.from_text(
                    text=f'[File: {doc["title"]}]\n' + '\n---\n'.join(chunks)
                ))
            loaded.append(doc['title'])
        except HTTPException:
            raise
        except Exception:
            raise HTTPException(502, f'Error reading "{doc.get("title", "document")}"')

    if not content_parts:
        raise HTTPException(400, 'Could not read any of the selected documents')

    try:
        answer = await gemini.ask_documents(content_parts, car, body.question)
        return {'answer': answer, 'sources': loaded}
    except Exception:
        raise HTTPException(502, 'AI document search failed — try again shortly')
